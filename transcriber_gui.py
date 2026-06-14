import os
import time
import threading
import subprocess
import tkinter as tk
from tkinter import filedialog, messagebox, ttk


def check_ffmpeg():
    try:
        subprocess.run(["ffmpeg", "-version"], capture_output=True, check=True)
        return True
    except (subprocess.CalledProcessError, FileNotFoundError):
        return False


class TranscriberGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Транскрибатор аудио/видео")
        self.root.geometry("680x520")
        self.root.minsize(600, 400)

        self.model = None
        self.is_transcribing = False

        self._create_widgets()
        self._log("Программа готова к работе.")
        self._log("Поддерживаемые форматы: MP3, WAV, FLAC, OGG, M4A, MP4, AVI, MKV, MOV и др.")
        if not check_ffmpeg():
            self._log("ВНИМАНИЕ: ffmpeg не найден. Установите ffmpeg для обработки видео.")

    def _create_widgets(self):
        main_frame = ttk.Frame(self.root, padding=10)
        main_frame.pack(fill=tk.BOTH, expand=True)

        # File selection
        ttk.Label(main_frame, text="Файл:").grid(row=0, column=0, sticky="w", pady=(0, 5))
        self.file_entry = ttk.Entry(main_frame)
        self.file_entry.grid(row=0, column=1, sticky="ew", padx=(5, 5), pady=(0, 5))
        ttk.Button(main_frame, text="Обзор...", command=self._browse_file).grid(
            row=0, column=2, padx=(0, 0), pady=(0, 5)
        )

        # Model selection
        ttk.Label(main_frame, text="Модель:").grid(row=1, column=0, sticky="w", pady=(0, 5))
        self.model_var = tk.StringVar(value="base")
        model_combo = ttk.Combobox(
            main_frame,
            textvariable=self.model_var,
            values=["tiny", "base", "small", "medium", "large"],
            state="readonly",
            width=12,
        )
        model_combo.grid(row=1, column=1, sticky="w", padx=(5, 5), pady=(0, 5))
        ttk.Label(
            main_frame,
            text="(tiny — быстро, large — точнее)",
            foreground="gray",
        ).grid(row=1, column=2, sticky="w", pady=(0, 5))

        # Output path
        ttk.Label(main_frame, text="Сохранить как:").grid(row=2, column=0, sticky="w", pady=(0, 5))
        self.output_entry = ttk.Entry(main_frame)
        self.output_entry.grid(row=2, column=1, sticky="ew", padx=(5, 5), pady=(0, 5))
        ttk.Button(main_frame, text="Обзор...", command=self._browse_output).grid(
            row=2, column=2, padx=(0, 0), pady=(0, 5)
        )

        # Start button
        btn_frame = ttk.Frame(main_frame)
        btn_frame.grid(row=3, column=0, columnspan=3, pady=10)

        self.start_button = tk.Button(
            btn_frame,
            text="Начать транскрибацию",
            command=self._start_transcription,
            bg="#4CAF50",
            fg="white",
            font=("Segoe UI", 10, "bold"),
            padx=20,
            pady=5,
            cursor="hand2",
        )
        self.start_button.pack()

        # Progress bar
        self.progress = ttk.Progressbar(main_frame, mode="indeterminate", length=300)
        self.progress.grid(row=4, column=0, columnspan=3, sticky="ew", padx=5, pady=(0, 5))

        # Status
        self.status_label = ttk.Label(main_frame, text="Готов к работе", foreground="blue")
        self.status_label.grid(row=5, column=0, columnspan=3, pady=(0, 5))

        # Log area
        ttk.Label(main_frame, text="Лог:").grid(row=6, column=0, sticky="nw", pady=(0, 2))
        log_frame = ttk.Frame(main_frame)
        log_frame.grid(row=7, column=0, columnspan=3, sticky="nsew", pady=(0, 0))

        self.log_text = tk.Text(log_frame, height=12, wrap=tk.WORD, state="disabled", font=("Consolas", 9))
        self.log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        scrollbar = ttk.Scrollbar(log_frame, command=self.log_text.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.log_text.configure(yscrollcommand=scrollbar.set)

        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(7, weight=1)

    def _browse_file(self):
        path = filedialog.askopenfilename(
            title="Выберите аудио или видео файл",
            filetypes=[
                ("Аудио/Видео", "*.mp3 *.wav *.flac *.ogg *.m4a *.mp4 *.avi *.mkv *.mov *.wmv *.webm"),
                ("Аудио", "*.mp3 *.wav *.flac *.ogg *.m4a"),
                ("Видео", "*.mp4 *.avi *.mkv *.mov *.wmv *.webm"),
                ("Все файлы", "*.*"),
            ],
        )
        if not path:
            return
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, path)
        base, _ = os.path.splitext(path)
        self.output_entry.delete(0, tk.END)
        self.output_entry.insert(0, base + "_транскрипция.docx")

    def _browse_output(self):
        path = filedialog.asksaveasfilename(
            title="Сохранить транскрипцию как",
            defaultextension=".docx",
            filetypes=[("Документ Word", "*.docx")],
        )
        if not path:
            return
        self.output_entry.delete(0, tk.END)
        self.output_entry.insert(0, path)

    def _log(self, msg):
        self.log_text.configure(state="normal")
        self.log_text.insert(tk.END, f"[{time.strftime('%H:%M:%S')}] {msg}\n")
        self.log_text.see(tk.END)
        self.log_text.configure(state="disabled")
        self.root.update_idletasks()

    def _start_transcription(self):
        if self.is_transcribing:
            return

        file_path = self.file_entry.get().strip()
        if not file_path or not os.path.exists(file_path):
            messagebox.showerror("Ошибка", "Выберите существующий файл")
            return

        output_path = self.output_entry.get().strip()
        if not output_path:
            base, _ = os.path.splitext(file_path)
            output_path = base + "_транскрипция.docx"
            self.output_entry.insert(0, output_path)

        self.is_transcribing = True
        self.start_button.configure(state="disabled", text="Идёт транскрибация...")
        self.progress.start()
        self.status_label.configure(text="Идёт транскрибация...", foreground="orange")

        thread = threading.Thread(
            target=self._transcribe_worker, args=(file_path, output_path), daemon=True
        )
        thread.start()

    def _transcribe_worker(self, file_path, output_path):
        try:
            import os as _os
            _os.environ["OMP_NUM_THREADS"] = "1"
            _os.environ["MKL_NUM_THREADS"] = "1"

            from faster_whisper import WhisperModel

            model_name = self.model_var.get()
            self._log(f"Загрузка модели faster-whisper ({model_name})...")
            self._log("Это может занять некоторое время при первом запуске (скачивание модели).")

            model = WhisperModel(
                model_name, device="cpu", compute_type="int8", cpu_threads=1, num_workers=1
            )

            self._log(f"Модель '{model_name}' загружена. Начинаю распознавание...")
            self._log("Это может занять продолжительное время в зависимости от размера файла.")

            segments, info = model.transcribe(file_path, language="ru")

            segments_list = list(segments)
            self._log(f"Распознавание завершено. Найдено {len(segments_list)} сегментов.")
            self._log("Создание документа Word...")

            self._save_to_docx(segments_list, output_path, model_name)

            file_size = os.path.getsize(output_path)
            size_str = self._format_size(file_size)
            self._log(f"Готово! Документ сохранён: {output_path} ({size_str})")
            self.root.after(0, lambda: self._on_complete(True, output_path))

        except ImportError:
            self._log("ОШИБКА: Библиотека faster-whisper не установлена.")
            self._log("Установите: pip install faster-whisper")
            self.root.after(0, lambda: self._on_complete(False, "faster-whisper не установлен"))
        except RuntimeError as e:
            msg = str(e)
            self._log(f"ОШИБКА: {msg}")
            if "mkl_malloc" in msg.lower() or "memory" in msg.lower():
                self._log("Недостаточно памяти. Попробуйте модель 'tiny' или 'base'.")
                self._log("Закройте другие программы и повторите попытку.")
            self.root.after(0, lambda: self._on_complete(False, msg))
        except Exception as e:
            self._log(f"ОШИБКА: {str(e)}")
            self.root.after(0, lambda: self._on_complete(False, str(e)))

    def _save_to_docx(self, segments, output_path, model_name):
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        section = doc.sections[0]
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

        title = doc.add_heading("Транскрипция", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta.add_run(f"Файл: {os.path.basename(self.file_entry.get())}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        meta.add_run("\n")
        run = meta.add_run(f"Модель: faster-whisper ({model_name}) | Язык: Русский")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        meta.add_run("\n")
        run = meta.add_run(
            f"Создано: {time.strftime('%d.%m.%Y %H:%M')} | "
            f"Сегментов: {len(segments)}"
        )
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()

        for segment in segments:
            start_time = segment.start
            text = segment.text.strip()
            if not text:
                continue

            hours = int(start_time // 3600)
            minutes = int((start_time % 3600) // 60)
            seconds = int(start_time % 60)
            timestamp = f"[{hours:02d}:{minutes:02d}:{seconds:02d}]"

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.3

            run = p.add_run(timestamp + " ")
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Consolas"

            run = p.add_run(text)
            run.font.size = Pt(12)

        doc.save(output_path)

    @staticmethod
    def _format_size(size_bytes):
        for unit in ["Б", "КБ", "МБ"]:
            if size_bytes < 1024:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024
        return f"{size_bytes:.1f} ГБ"

    def _on_complete(self, success, message):
        self.progress.stop()
        self.start_button.configure(state="normal", text="Начать транскрибацию")
        if success:
            self.status_label.configure(text="Готово! Документ создан.", foreground="green")
            messagebox.showinfo("Готово", f"Транскрипция сохранена:\n{message}")
        else:
            self.status_label.configure(text=f"Ошибка", foreground="red")
            messagebox.showerror("Ошибка", message)
        self.is_transcribing = False


def main():
    root = tk.Tk()
    app = TranscriberGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()
